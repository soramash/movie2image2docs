#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import argparse
import concurrent.futures
import logging
import os
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

# OpenAI SDK
try:
    from openai import OpenAI
except Exception:
    OpenAI = None  # type: ignore

# python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except Exception:
    Document = None  # type: ignore

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}


@dataclass
class OcrResult:
    image_path: Path
    markdown_text: str


def ensure_dir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)


def find_section_dirs(root: Path, pattern: str) -> List[Path]:
    return sorted([p for p in root.glob(pattern) if p.is_dir()])


def list_images(section_dir: Path) -> List[Path]:
    return [p for p in sorted(section_dir.iterdir()) if p.suffix.lower() in IMAGE_EXTS]


def load_cached_md(cache_dir: Path, section: str, image_name: str) -> Optional[str]:
    cache_path = cache_dir / section / f"{image_name}.md"
    if cache_path.exists():
        try:
            return cache_path.read_text(encoding="utf-8")
        except Exception:
            return None
    return None


def save_cached_md(cache_dir: Path, section: str, image_name: str, md: str) -> None:
    cache_path = cache_dir / section / f"{image_name}.md"
    ensure_dir(cache_path.parent)
    cache_path.write_text(md, encoding="utf-8")


def ocr_markdown_with_openai(image_path: Path, model: str, language_hint: str = "ja") -> str:
    if OpenAI is None:
        raise RuntimeError("openai SDK is not installed. Install dependencies from requirements-docx.txt")
    if not os.environ.get("OPENAI_API_KEY"):
        raise RuntimeError("OPENAI_API_KEY is not set. Export it before running.")

    client = OpenAI()

    # Use Chat Completions with image input via image_url (file path must be converted to data URL or file URL)
    # Simpler: use file:// URL; however OpenAI API expects hosted or data URL. We'll use data URL.
    import base64
    mime = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".bmp": "image/bmp",
    }.get(image_path.suffix.lower(), "application/octet-stream")
    b64 = base64.b64encode(image_path.read_bytes()).decode("ascii")
    data_url = f"data:{mime};base64,{b64}"

    prompt = (
        "次の画像から可読な文字情報を正確にMarkdownで抽出してください。\n"
        "要件:\n"
        "- 出力は有効なMarkdownのみ（前置き/後置きの説明は禁止）\n"
        "- 見出し/強調/箇条書き/番号リスト/表は可能な限りMarkdownで再現\n"
        "- 不要なノイズや透かしは除外\n"
        f"- 言語は原文のまま（{language_hint} を優先）\n"
        "- 画像に文字がない場合は (テキストなし) とだけ出力\n"
    )

    try:
        resp = client.chat.completions.create(
            model=model,
            temperature=0,
            messages=[
                {"role": "system", "content": "You are a precise OCR that outputs only Markdown."},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                },
            ],
        )
        text = resp.choices[0].message.content or ""
        return text.strip()
    except Exception as e:
        raise RuntimeError(f"OpenAI OCR failed for {image_path}: {e}")


def preflight_openai(model: str, timeout: float = 15.0) -> None:
    """Fail-fast connectivity/auth check against OpenAI API and model availability."""
    if OpenAI is None:
        raise RuntimeError("openai SDK is not installed. Install dependencies from requirements-docx.txt")
    if not os.environ.get("OPENAI_API_KEY"):
        raise RuntimeError("OPENAI_API_KEY is not set. Export it before running.")
    try:
        client = OpenAI(timeout=timeout)
        models = client.models.list()
        ids = {getattr(m, "id", None) for m in getattr(models, "data", []) if getattr(m, "id", None)}
        # If model is specified and not present, fail early with a helpful message
        if model and ids and model not in ids:
            examples = ", ".join(sorted(list(ids))[:10])
            raise RuntimeError(f"指定モデルが見つかりません: {model}. 利用可能なモデル例: {examples}")
        logging.info("OpenAI preflight OK")
    except Exception as e:
        raise RuntimeError(f"OpenAIへの疎通に失敗しました（認証/ネットワーク/権限）: {e}")


def process_section(
    section_dir: Path,
    model: str,
    cache_dir: Path,
    language_hint: str,
    max_workers: int,
) -> Tuple[str, List[OcrResult]]:
    images = list_images(section_dir)
    if not images:
        logging.warning(f"画像が見つかりません: {section_dir}")
        return (section_dir.name, [])

    results: List[OcrResult] = []

    def _one(img_path: Path) -> OcrResult:
        cached = load_cached_md(cache_dir, section_dir.name, img_path.name)
        if cached is None:
            md = ocr_markdown_with_openai(img_path, model=model, language_hint=language_hint)
            save_cached_md(cache_dir, section_dir.name, img_path.name, md)
        else:
            md = cached
        return OcrResult(image_path=img_path, markdown_text=md)

    if max_workers > 1:
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as ex:
            for r in ex.map(_one, images):
                results.append(r)
    else:
        for img in images:
            results.append(_one(img))

    return (section_dir.name, results)


def _ensure_markdown_style(doc) -> str:
    """Create a monospace paragraph style for raw Markdown text if not exists."""
    style_name = "Markdown"
    styles = doc.styles
    try:
        styles[style_name]
        return style_name
    except KeyError:
        pass
    style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Courier New"
    font.size = Pt(10)
    pf = style.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.15
    return style_name


def _page_content_width_inches(doc) -> float:
    sec = doc.sections[0]
    # EMU -> inches (1 inch = 914400 EMU)
    width_emu = sec.page_width - sec.left_margin - sec.right_margin
    return float(width_emu) / 914400.0


def export_section_to_docx(title: str, results: List[OcrResult], out_path: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Install dependencies from requirements-docx.txt")

    doc = Document()
    doc.add_heading(title, level=0)

    style_name = _ensure_markdown_style(doc)
    max_width_in = _page_content_width_inches(doc)

    for idx, r in enumerate(results, 1):
        doc.add_heading(f"{idx:02d}. {r.image_path.name}", level=1)
        doc.add_paragraph(str(r.image_path))
        # Insert image scaled to page content width
        p = doc.add_paragraph()
        run = p.add_run()
        try:
            run.add_picture(str(r.image_path), width=Inches(max_width_in))
        except Exception:
            # Fallback without scaling
            run.add_picture(str(r.image_path))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("抽出テキスト（Markdown原文）", level=2)
        # Add a single paragraph with explicit line breaks to preserve Markdown layout
        lines = r.markdown_text.splitlines()
        para = doc.add_paragraph(style=style_name)
        if not lines:
            lines = [""]
        for i, line in enumerate(lines):
            run = para.add_run(line)
            if i < len(lines) - 1:
                run.add_break()

    ensure_dir(out_path.parent)
    doc.save(str(out_path))


def sanitize_title(name: str, prefix: Optional[str] = None) -> str:
    import re
    title = re.sub(r"[\\/:*?\"<>|]", "_", name)
    if prefix:
        title = f"{prefix} {title}"
    return title


def main() -> None:
    parser = argparse.ArgumentParser(description="セクション画像からWord(.docx)ドキュメントを生成（OpenAI OCRでMarkdown抽出）")
    parser.add_argument("--root", default=".", help="探索ルートディレクトリ（デフォルト: .）")
    parser.add_argument("--pattern", default="slides-vfr-IPA-20250930_section0*", help="セクションフォルダのglobパターン")
    parser.add_argument("--model", default="gpt-4o-mini", help="OpenAIのビジョン対応モデル名（例: gpt-4o-mini）")
    parser.add_argument("--language", default="ja", help="OCRの言語ヒント（ja/enなど）")
    parser.add_argument("--max-workers", type=int, default=2, help="並列OCR数（初期値2）")
    parser.add_argument("--title-prefix", default=None, help="ドキュメントタイトルの先頭に付ける文字列")
    parser.add_argument("--cache-dir", default=".cache/ocr-md", help="OCR結果キャッシュ保存先")
    parser.add_argument("--out-dir", default="build/word", help="出力ディレクトリ（.docx）")
    parser.add_argument("--log-level", default="INFO", help="ログレベル（DEBUG/INFO/WARN/ERROR）")
    parser.add_argument("--no-preflight", action="store_true", help="OpenAIへの疎通チェックをスキップ")

    args = parser.parse_args()

    logging.basicConfig(level=getattr(logging, args.log_level.upper(), logging.INFO), format="%(levelname)s %(message)s")

    root = Path(args.root).resolve()
    cache_dir = Path(args.cache_dir)
    out_dir = Path(args.out_dir)
    ensure_dir(cache_dir)
    ensure_dir(out_dir)

    if not args.no_preflight:
        try:
            preflight_openai(args.model)
        except Exception as e:
            logging.error(str(e))
            raise SystemExit(2)

    sections = find_section_dirs(root, args.pattern)
    if not sections:
        logging.error(f"セクションフォルダが見つかりません: root={root}, pattern={args.pattern}")
        raise SystemExit(1)

    summary: List[Tuple[str, Path]] = []

    for sec in sections:
        sec_name, results = process_section(
            section_dir=sec,
            model=args.model,
            cache_dir=cache_dir,
            language_hint=args.language,
            max_workers=args.max_workers,
        )
        title = sanitize_title(sec_name, args.title_prefix)
        out_path = out_dir / f"{title}.docx"
        export_section_to_docx(title, results, out_path)
        logging.info(f"Wordを出力しました: {out_path}")
        summary.append((title, out_path))

    print("\n=== 生成結果(.docx) ===")
    for title, path in summary:
        print(f"- {title}: {path}")


if __name__ == "__main__":
    main()
